import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_inline_formatting(paragraph, text: str):
    """Parse inline markdown like **bold** and *italic* and add runs to paragraph."""
    # A simple regex to split out **bold** and *italic*
    # We will split by both to preserve order
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def markdown_to_docx(markdown_text: str, output_path: str | Path) -> None:
    """Safely convert basic markdown to a nicely formatted docx."""
    doc = Document()

    # Apply some basic document styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            p = doc.add_heading(level=1)
            apply_inline_formatting(p, line[2:].strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            apply_inline_formatting(p, line[3:].strip())
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            apply_inline_formatting(p, line[4:].strip())
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, line[2:].strip())
        elif line == '---' or line == '***':
            # Skip horizontal rules for now, or add an empty line
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            apply_inline_formatting(p, line)

    doc.save(str(output_path))
